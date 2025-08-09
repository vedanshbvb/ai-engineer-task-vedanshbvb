import os
import mimetypes
from urllib.parse import urljoin
from bs4 import BeautifulSoup
import requests
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry

from langchain_community.document_loaders import WebBaseLoader, PyPDFLoader
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document as LC_Document

from docx import Document as DocxDocument
from PyPDF2 import PdfReader
import pymupdf as fitz  # PyMuPDF

# ---------------- PATHS ----------------
DATA_SOURCES_PDF = "data/given/Data Sources.pdf"
WEBPAGE_DIR = "data/raw/webpages"
TEMPLATE_DIR = "data/raw/templates"
VECTORSTORE_DIR = "data/vectorstore"

# ---------------- SESSION ----------------
DEFAULT_USER_AGENT = os.environ.get(
    "USER_AGENT",
    "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 "
    "(KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
)

session = requests.Session()
retries = Retry(total=3, backoff_factor=1, status_forcelist=[429, 500, 502, 503, 504])
session.mount("https://", HTTPAdapter(max_retries=retries))
session.mount("http://", HTTPAdapter(max_retries=retries))

# ---------------- UTILS ----------------
def ensure_dirs():
    os.makedirs(WEBPAGE_DIR, exist_ok=True)
    os.makedirs(TEMPLATE_DIR, exist_ok=True)
    os.makedirs(VECTORSTORE_DIR, exist_ok=True)

def extract_links_from_pdf(pdf_path):
    """Extract all hyperlinks from a PDF file using PyMuPDF."""
    links = []
    doc = fitz.open(pdf_path)
    for page in doc:
        for link in page.get_links():
            uri = link.get("uri")
            if uri:
                links.append(uri)
    return links

def _safe_filename_from_cd(cd):
    if not cd:
        return None
    import re
    m = re.search(r'filename\*=UTF-8\'\'(?P<f>[^;]+)', cd)
    if m:
        return m.group("f")
    m = re.search(r'filename="?([^";]+)"?', cd)
    if m:
        return m.group(1)
    return None

def download_file(url, dest_dir, referer=None, skip_if_exists=True):
    """Robust download: handles redirects, content-disposition, HTML landing pages."""
    os.makedirs(dest_dir, exist_ok=True)
    headers = {
        "User-Agent": DEFAULT_USER_AGENT,
        "Accept": "*/*",
    }
    if referer:
        headers["Referer"] = referer

    try:
        resp = session.get(url, headers=headers, stream=True, allow_redirects=True, timeout=30)
        status = resp.status_code

        if status != 200:
            print(f"[!] GET {url} returned status {status}")
            content_type = resp.headers.get("content-type", "")
            if "text/html" in content_type.lower():
                html = resp.text
                soup = BeautifulSoup(html, "html.parser")
                for a in soup.find_all("a", href=True):
                    href = a["href"]
                    if href.lower().endswith((".docx", ".pdf")) or "download" in href.lower():
                        file_url = urljoin(url, href)
                        print(f"[+] Found downloadable link in landing page: {file_url}")
                        return download_file(file_url, dest_dir, referer=url, skip_if_exists=skip_if_exists)
            return None

        content_disp = resp.headers.get("content-disposition")
        filename = _safe_filename_from_cd(content_disp)
        if not filename:
            filename = url.split("/")[-1] or "downloaded_file"
        filename = filename.split("?")[0].split("#")[0]
        dest_path = os.path.join(dest_dir, filename)

        if skip_if_exists and os.path.exists(dest_path):
            print(f"[=] Already downloaded (skip): {filename}")
            return dest_path

        with open(dest_path, "wb") as fh:
            for chunk in resp.iter_content(chunk_size=8192):
                if chunk:
                    fh.write(chunk)

        print(f"[+] Downloaded: {dest_path}")
        return dest_path

    except requests.exceptions.RequestException as e:
        print(f"[!] Request failed for {url}: {e}")
        return None

def scrape_webpage(url):
    """Scrape webpage text and download any linked .docx/.pdf."""
    print(f"[+] Scraping webpage: {url}")
    fname = url.split("/")[-1] or "index"
    fname = fname.replace(".html", "").replace("/", "_") or "index"
    text_path = os.path.join(WEBPAGE_DIR, f"{fname}.txt")

    if not os.path.exists(text_path):
        try:
            loader = WebBaseLoader(url)
            docs = loader.load()
            page_text = docs[0].page_content
            with open(text_path, "w", encoding="utf-8") as f:
                f.write(page_text)
        except Exception as e:
            print(f"[!] Failed WebBaseLoader for {url}: {e}")
            try:
                resp = session.get(url, headers={"User-Agent": DEFAULT_USER_AGENT}, timeout=20)
                with open(text_path, "w", encoding="utf-8") as f:
                    f.write(resp.text)
            except Exception as e2:
                print(f"[!] Could not fetch page text for {url}: {e2}")

    # Find and download docs linked on the page
    try:
        html = session.get(url, headers={"User-Agent": DEFAULT_USER_AGENT}, timeout=20).text
        soup = BeautifulSoup(html, "html.parser")
        for a in soup.find_all("a", href=True):
            href = a["href"]
            if href.lower().endswith((".docx", ".pdf")) or "download" in href.lower() or "assets" in href.lower():
                file_url = urljoin(url, href)
                download_file(file_url, TEMPLATE_DIR, referer=url, skip_if_exists=True)
    except Exception as e:
        print(f"[!] Error parsing links from {url}: {e}")

def extract_text_from_docx(path):
    doc = DocxDocument(path)
    return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

# ---------------- INGESTION ----------------
def ingest_all():
    ensure_dirs()
    links = extract_links_from_pdf(DATA_SOURCES_PDF)
    print(f"[+] Found {len(links)} links in Data Sources.pdf")

    for link in links:
        mime_type, _ = mimetypes.guess_type(link)
        if link.lower().endswith((".docx", ".pdf")) or (mime_type and ("word" in mime_type or "pdf" in mime_type)):
            download_file(link, TEMPLATE_DIR)
        else:
            scrape_webpage(link)

    # Gather all text content
    texts = []

    # From webpages
    for fname in os.listdir(WEBPAGE_DIR):
        with open(os.path.join(WEBPAGE_DIR, fname), "r", encoding="utf-8") as f:
            texts.append(f.read())

    # From templates
    for fname in os.listdir(TEMPLATE_DIR):
        path = os.path.join(TEMPLATE_DIR, fname)
        if fname.lower().endswith(".docx"):
            texts.append(extract_text_from_docx(path))
        elif fname.lower().endswith(".pdf"):
            loader = PyPDFLoader(path)
            for doc in loader.load():
                texts.append(doc.page_content)

    # Chunk and embed
    splitter = RecursiveCharacterTextSplitter(chunk_size=800, chunk_overlap=100)
    chunks = []
    for text in texts:
        chunks.extend(splitter.split_text(text))

    print(f"[+] Total chunks: {len(chunks)}")
    embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
    docs = [LC_Document(page_content=c, metadata={}) for c in chunks]
    vectorstore = FAISS.from_documents(docs, embeddings)
    vectorstore.save_local(VECTORSTORE_DIR)
    print(f"[+] Saved FAISS index to {VECTORSTORE_DIR}")

if __name__ == "__main__":
    ingest_all()
