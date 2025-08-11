import os
from typing import List, Tuple
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from docx import Document as DocxDocument
import PyPDF2
from google import genai
from google.genai import types
from dotenv import load_dotenv
import json
from docx.shared import Pt
from docx.enum.text import WD_COLOR_INDEX
import re

load_dotenv()

VECTORSTORE_PATH = os.path.join("data", "vectorstore")

# ----------------- Load FAISS -----------------
def load_faiss_vectorstore():
    embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
    return FAISS.load_local(VECTORSTORE_PATH, embeddings, allow_dangerous_deserialization=True)

# ----------------- Extractors -----------------
def extract_text_from_docx(path) -> str:
    doc = DocxDocument(path)
    return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

def extract_text_from_pdf(path) -> str:
    text = ""
    with open(path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        for page in reader.pages:
            text += page.extract_text() or ""
    return text

# ----------------- Chunk Helper -----------------
def chunk_text(text: str, chunk_size=800, chunk_overlap=100):
    splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    return splitter.split_text(text)

# ----------------- Editing user docs -----------------
import re
from docx import Document as DocxDocument
from docx.shared import Pt
from docx.enum.text import WD_COLOR_INDEX




# def highlight_and_comment_docx(input_path: str, output_path: str, issues: list):
#     """
#     Improved: Highlights and adds inline comments to a DOCX file based on issues.
#     - Matches keywords flexibly (section, issue, placeholders from suggestion)
#     - Case-insensitive
#     - Falls back to positional matches for vague section names
#     """
#     if not issues:
#         return False

#     doc = DocxDocument(input_path)
#     updated = False

#     # Helper: highlight and comment text inside a paragraph
#     def apply_highlight(para, keyword, comment):
#         nonlocal updated
#         pattern = re.compile(re.escape(keyword), re.IGNORECASE)
#         if keyword.lower() in para.text.lower():
#             for run in para.runs:
#                 if pattern.search(run.text):
#                     run.font.highlight_color = WD_COLOR_INDEX.YELLOW
#                     comment_run = para.add_run(f" [COMMENT: {comment}]")
#                     comment_run.italic = True
#                     comment_run.font.size = Pt(10)
#                     updated = True

#     for issue in issues:
#         section = issue.get("section", "").strip()
#         issue_text = issue.get("issue", "").strip()
#         suggestion = issue.get("suggestion", "").strip()

#         # Build list of potential keywords to search for
#         keywords = []

#         # Add section name if it’s specific (not just "General")
#         if section and section.lower() not in ["general", "header", "signatures", "first paragraph"]:
#             keywords.append(section)

#         # Add placeholders from suggestion/issue if any
#         placeholder_match = re.findall(r"\{[^}]+\}", suggestion + " " + issue_text)
#         keywords.extend([ph for ph in placeholder_match])

#         # Add short words from issue text that might appear in doc
#         if not placeholder_match and issue_text:
#             words = [w for w in issue_text.split() if len(w) > 4]  # skip very short words
#             keywords.extend(words[:3])  # limit to first few

#         matched = False

#         # 1️⃣ Keyword-based highlighting
#         for para in doc.paragraphs:
#             for kw in keywords:
#                 if kw and kw.lower() in para.text.lower():
#                     apply_highlight(para, kw, suggestion)
#                     matched = True

#         # 2️⃣ Positional fallback for vague section labels
#         if not matched:
#             if section.lower() == "header" and doc.paragraphs:
#                 apply_highlight(doc.paragraphs[0], doc.paragraphs[0].text, suggestion)

#             elif section.lower() == "first paragraph" and len(doc.paragraphs) > 1:
#                 apply_highlight(doc.paragraphs[1], doc.paragraphs[1].text, suggestion)

#             elif section.lower() == "signatures":
#                 for para in doc.paragraphs[-5:]:  # look in last 5 paragraphs
#                     if "sign" in para.text.lower() or "signature" in para.text.lower():
#                         apply_highlight(para, para.text, suggestion)

#     # Always save, even if no highlights
#     doc.save(output_path)
#     return updated


def highlight_and_comment_docx(input_path: str, output_path: str, issues: list):
    """
    Highlights and adds inline comments to a DOCX file based on issues.
    - Exact match for section: "5. Governing Laws" OR "Governing Laws"
    - If section is 'General', just add comment at the end (no highlights)
    - Always saves the file even if no highlights
    """
    if not issues:
        doc = DocxDocument(input_path)
        doc.save(output_path)
        return False

    doc = DocxDocument(input_path)
    updated = False

    def apply_highlight(para, keyword, comment):
        """Highlight if exact match found, then add comment once."""
        nonlocal updated
        if keyword.lower() in para.text.lower():
            # Exact match check: full section or section without numbering
            section_plain = re.sub(r"^\d+\.\s*", "", keyword, flags=re.IGNORECASE).strip().lower()
            para_plain = para.text.lower()

            if section_plain in para_plain or keyword.lower() in para_plain:
                # Highlight
                for run in para.runs:
                    if section_plain in run.text.lower() or keyword.lower() in run.text.lower():
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                # Add comment if not already there
                if not any("[COMMENT:" in r.text for r in para.runs):
                    comment_run = para.add_run(f" [COMMENT: {comment}]")
                    comment_run.italic = True
                    comment_run.font.size = Pt(10)
                updated = True

    for issue in issues:
        section = issue.get("section", "").strip()
        suggestion = issue.get("suggestion", "").strip()

        # If section is "General" → append comment at the end
        if section.lower() == "general":
            end_para = doc.add_paragraph(f"[COMMENT: {suggestion}]")
            end_para.italic = True
            end_para.font.size = Pt(10)
            updated = True
            continue

        # Try to match full section (e.g., "5. Governing Laws") or without number
        section_no_number = re.sub(r"^\d+\.\s*", "", section, flags=re.IGNORECASE).strip()

        matched = False
        for para in doc.paragraphs:
            if section and (
                section.lower() in para.text.lower() or
                section_no_number.lower() in para.text.lower()
            ):
                apply_highlight(para, section, suggestion)
                matched = True

        # Positional fallback for signatures
        if not matched and section.lower() == "signatures":
            for para in doc.paragraphs[-5:]:
                if "sign" in para.text.lower() or "signature" in para.text.lower():
                    apply_highlight(para, para.text, suggestion)

    # Always save, even if nothing matched
    doc.save(output_path)
    return updated




# ----------------- Gemini Call -----------------
def call_gemini_combined(user_docs: str, references: str) -> str:
    client = genai.Client(api_key=os.environ.get("GEMINI_API_KEY"))
    model = "gemini-2.0-flash"

    prompt = f"""
    You are a compliance assistant. Compare the following user document chunk to the reference clauses below. 
    Identify what is the type of ADGM(Abu Dhabi Global Market) document the user has sent (for example, application form, mou) and also identify what is the user trying to do, for example, company formation, employment contract etc etc. What other documents does the user need to provide to complete the process? 

    You must:
    1. Identify the legal process they are trying to complete.
    2. Suggest fixes.
    3. Compare the uploaded documents with the required ADGM checklist for that process.
    4. Detect compliance issues, red flags, or ADGM violations.
    5. Identify the document uploaded, they will be one of the following:
        - Articles of Association
        - Memorandum of Association
        - Board Resolution
        - Shareholder Resolution
        - Incorporation Application Form
        - UBO Declaration form
        - Register of Members and Directors
        - Change of Registered Address Notice
    

    Identify any compliance issues, missing elements such as these:

    Red Flag Detection Features
        • Invalid or missing clauses
        • Incorrect jurisdiction (e.g., referencing UAE Federal Courts instead of ADGM)
        • Ambiguous or non-binding language like "today", "tomorrow", "maybe", "monday", "this week"
        • Missing signatory sections or improper formatting
        • Non-compliance with ADGM-specific templates
        • Missing documents
        • Unresolved placeholders, for example if the doc contains something like {{name}} or {{company}}
        

    So you must identify missing documents, issues in the user documents, their severity and your suggestion to fix it.

    Your answer must STRICTLY be in json format:
    {{
        "issues_found": [
            {{
                "document": "<string>",
                "section": "<string or null>",
                "issue": "<string>",
                "severity": "<Low/Medium/High>",
                "suggestion": "<string>"
            }}
        ]
    }}



    For example:
        {{
            "issues_found": 
            [
                {{
                    "document": "UBO Declaration form",
                    "section": "1. Governing Laws",
                    "issue": "Jurisdiction clause does not specify ADGM",
                    "severity": "High",
                    "suggestion": "Update jurisdiction to ADGM Courts."
                }}
            ]
        }}

    The section should be a number followed by the name of the section
    NO value should ever be Null!

    User Documents:
    {user_docs}

    Reference Clauses:
    {references}


    """




    contents = [types.Content(role="user", parts=[types.Part.from_text(text=prompt)])]
    # print(references)

    response_text = ""
    for chunk in client.models.generate_content_stream(model=model, contents=contents):
        if chunk.text:
            response_text += chunk.text
    return response_text.strip()

# ----------------- Main Review -----------------
def review_documents(filepaths: List[str]) -> dict:
    vectorstore = load_faiss_vectorstore()

    # List of all required docs for Company Incorporation
    required_docs = [
        "Articles of Association",
        "Memorandum of Association",
        "Board Resolution",
        "Shareholder Resolution",
        "Incorporation Application Form",
        "UBO Declaration form",
        "Register of Members and Directors",
        "Change of Registered Address Notice"
    ]

    all_issues = {"issues_found": []}
    uploaded_doc_names = []

    # ---- Count uploaded docs ----
    all_issues["process"] = "Company Incorporation"
    all_issues["documents_uploaded"] = len(uploaded_doc_names)
    all_issues["required_documents"] = len(required_docs)


    for path in filepaths:
        ext = os.path.splitext(path)[1].lower()
        if ext == ".docx":
            text = extract_text_from_docx(path)
        elif ext == ".pdf":
            text = extract_text_from_pdf(path)
        else:
            continue

        fname = os.path.basename(path)
        uploaded_doc_names.append(fname)

        gemini_text = f"\n### Document: {fname}\n{text}\n"

        # ---- RAG retrieval using ALL chunks ----
        ref_texts = []
        chunks = chunk_text(text)
        for chunk in chunks:
            docs = vectorstore.similarity_search(chunk, k=3)
            ref_texts.extend([doc.page_content for doc in docs])
        references_combined = "\n---\n".join(list(set(ref_texts)))

        # ---- Call Gemini for THIS document ----
        issues_json = call_gemini_combined(gemini_text, references_combined)

        # ---- Clean and parse Gemini output ----
        import re
        try:
            cleaned = issues_json.strip()

            # If wrapped in ```json ... ``` or ``` ... ```
            if cleaned.startswith("```"):
                cleaned = re.sub(r"^```[a-zA-Z]*\n?", "", cleaned)
                cleaned = re.sub(r"\n?```$", "", cleaned)

            parsed = json.loads(cleaned.strip())
        except json.JSONDecodeError:
            parsed = {"error": "Invalid JSON response from Gemini.", "raw_output": issues_json}

        # ---- Merge issues into all_issues ----
        if isinstance(parsed, dict) and "issues_found" in parsed:
            all_issues["issues_found"].extend(parsed["issues_found"])


    # ---- Update uploaded docs ----
    all_issues["documents_uploaded"] = len(uploaded_doc_names)


    # ---- Determine missing documents ----
    uploaded_doc_types = {issue.get("document") for issue in all_issues["issues_found"] if "document" in issue}
    missing_docs = [doc for doc in required_docs if doc not in uploaded_doc_types]
    all_issues["missing_document"] = missing_docs

    return all_issues



if __name__ == "__main__":
    files = ["sample1.docx", "sample2.pdf"]  # replace with actual uploads
    output = review_documents(files)
    print(output)
