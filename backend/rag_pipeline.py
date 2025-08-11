import os
from typing import List, Tuple
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from docx import Document as DocxDocument
import PyPDF2
from google import genai
from google.genai import types
from dotenv import load_dotenv
import json
from docx.shared import Pt
from docx.enum.text import WD_COLOR_INDEX
import re

load_dotenv()

VECTORSTORE_PATH = os.path.join("data", "vectorstore")

# ----------------- Load FAISS -----------------
def load_faiss_vectorstore():
    embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
    return FAISS.load_local(VECTORSTORE_PATH, embeddings, allow_dangerous_deserialization=True)

# ----------------- Extractors -----------------
def extract_text_from_docx(path) -> str:
    doc = DocxDocument(path)
    return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

def extract_text_from_pdf(path) -> str:
    text = ""
    with open(path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        for page in reader.pages:
            text += page.extract_text() or ""
    return text

# ----------------- Chunk Helper -----------------
def chunk_text(text: str, chunk_size=800, chunk_overlap=100):
    splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    return splitter.split_text(text)

# ----------------- Editing user docs -----------------
def highlight_and_comment_docx(input_path: str, output_path: str, issues: list):
    """
    Highlights and adds inline comments to a DOCX file wherever issues are found.
    Saves the updated document to output_path.
    Matching is case-insensitive and allows partial matches.
    """
    if not issues:
        return False

    doc = DocxDocument(input_path)
    updated = False

    for issue in issues:
        keyword = issue.get("section") or issue.get("issue")
        if not keyword:
            continue

        # Build comment text
        comment_text = f"[COMMENT: {issue.get('suggestion', '')}]"

        # Case-insensitive keyword search in paragraphs
        pattern = re.compile(re.escape(keyword), re.IGNORECASE)

        for para in doc.paragraphs:
            if pattern.search(para.text):
                for run in para.runs:
                    if pattern.search(run.text):
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                        comment_run = para.add_run(f" {comment_text}")
                        comment_run.italic = True
                        comment_run.font.size = Pt(10)
                        updated = True

    if updated:
        doc.save(output_path)
    return updated




# ----------------- Gemini Call -----------------
def call_gemini_combined(user_docs: str, references: str) -> str:
    client = genai.Client(api_key=os.environ.get("GEMINI_API_KEY"))
    model = "gemini-2.0-flash"

    prompt = f"""
    You are a compliance assistant. Compare the following user document chunk to the reference clauses below. 
    Identify what is the type of ADGM(Abu Dhabi Global Market) document the user has sent (for example, application form, mou) and also identify what is the user trying to do, for example, company formation, employment contract etc etc. What other documents does the user need to provide to complete the process? 

    You must:
    1. Identify the legal process they are trying to complete.
    2. Identify the type of each document uploaded.
    3. Compare the uploaded documents with the required ADGM checklist for that process.
    4. List missing documents, if any.
    5. Detect compliance issues, red flags, or ADGM violations.
    6. Suggest fixes.

    
    For example, if a user wants to form a company, they will need the following documents:
    1. Articles of Association
    2. Memorandum of Association
    3. Board Resolution
    4. Shareholder Resolution
    5. Incorporation Application Form
    6. UBO Declaration form
    7. Register of Members and Directors
    8. Change of Registered Address Notice
    

    Identify any compliance issues, missing elements or documents such as these:

    Red Flag Detection Features
        • Invalid or missing clauses
        • Incorrect jurisdiction (e.g., referencing UAE Federal Courts instead of ADGM)
        • Ambiguous or non-binding language like "today", "tomorrow", "maybe" 
        • Missing signatory sections or improper formatting
        • Non-compliance with ADGM-specific templates
        • Missing documents
        • Unresolved placeholders, for example if the doc contains something like {{name}} or {{company}}
        

    So you must identify missing documents, issues in the user documents, their severity and your suggestion to fix it.

    Your answer must STRICTLY be in json format:
    {{
        "process": "<string>",
        "documents_uploaded": <int>,
        "required_documents": <int>,
        "missing_document": "<string or list>",
        "issues_found": [
            {{
                "document": "<string>",
                "section": "<string or null>",
                "issue": "<string>",
                "severity": "<Low/Medium/High>",
                "suggestion": "<string>"
            }}
        ]
    }}



    For example:
        {{
            "process": "Company Incorporation",
            "documents_uploaded": 4,
            "required_documents": 5,
            "missing_document": "Register of Members and Directors",
            "issues_found": 
            [
                {{
                    "document": "Articles of Association",
                    "section": "Clause 3.1",
                    "issue": "Jurisdiction clause does not specify ADGM",
                    "severity": "High",
                    "suggestion": "Update jurisdiction to ADGM Courts."
                }}
            ]
        }}

    User Documents:
    {user_docs}

    Reference Clauses:
    {references}


    """




    contents = [types.Content(role="user", parts=[types.Part.from_text(text=prompt)])]
    # print(references)

    response_text = ""
    for chunk in client.models.generate_content_stream(model=model, contents=contents):
        if chunk.text:
            response_text += chunk.text
    return response_text.strip()

# ----------------- Main Review -----------------
def review_documents(filepaths: List[str]) -> str:
    vectorstore = load_faiss_vectorstore()
    all_texts: List[Tuple[str, str]] = []

    # Extract all docs fully
    for path in filepaths:
        ext = os.path.splitext(path)[1].lower()
        if ext == ".docx":
            text = extract_text_from_docx(path)
        elif ext == ".pdf":
            text = extract_text_from_pdf(path)
        else:
            continue
        all_texts.append((os.path.basename(path), text))

    if not all_texts:
        return "No supported files uploaded."

    # Build combined text for Gemini
    combined_text = ""
    for fname, text in all_texts:
        combined_text += f"\n### Document: {fname}\n{text}\n"

    # RAG: collect relevant references using first chunk of each doc
    ref_texts = []
    for _, text in all_texts:
        chunks = chunk_text(text)
        if chunks:
            docs = vectorstore.similarity_search(chunks[0], k=3)
            ref_texts.extend([doc.page_content for doc in docs])
    references_combined = "\n---\n".join(list(set(ref_texts)))  # deduplicate

    # Single Gemini call
    # print(references_combined)
    issues_json = call_gemini_combined(combined_text, references_combined)

    # return issues_json
    import re
    try:
        cleaned = issues_json.strip()

        # If wrapped in ```json ... ``` or ``` ... ```
        if cleaned.startswith("```"):
            # Remove starting triple backticks and optional 'json'
            cleaned = re.sub(r"^```[a-zA-Z]*\n?", "", cleaned)
            # Remove trailing triple backticks
            cleaned = re.sub(r"\n?```$", "", cleaned)

        return json.loads(cleaned.strip())

    except json.JSONDecodeError:
        return {"error": "Invalid JSON response from Gemini.", "raw_output": issues_json}





if __name__ == "__main__":
    files = ["sample1.docx", "sample2.pdf"]  # replace with actual uploads
    output = review_documents(files)
    print(output)
